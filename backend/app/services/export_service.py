import json
from pathlib import Path
from datetime import datetime
from app.models.project import Project
from app.config import settings
from app.utils.logger import get_logger

logger = get_logger(__name__)


class ExportService:
    def _get_export_dir(self, project_id: int) -> Path:
        d = settings.exports_dir / str(project_id)
        d.mkdir(parents=True, exist_ok=True)
        return d

    def export_txt(self, project: Project) -> Path:
        out_dir = self._get_export_dir(project.id)
        path = out_dir / f"{project.id}_export.txt"

        lines = [f"PROJECT: {project.name}", f"Date: {project.created_at.strftime('%Y-%m-%d %H:%M')}", "=" * 60, ""]

        if project.transcription and project.transcription.text:
            lines += ["TRANSCRIPTION", "-" * 40, project.transcription.text, ""]

        for analysis in project.analyses:
            if analysis.content:
                lines += [f"\n{analysis.analysis_type.upper().replace('_', ' ')}", "-" * 40, analysis.content, ""]

        path.write_text("\n".join(lines), encoding="utf-8")
        return path

    def export_markdown(self, project: Project) -> Path:
        out_dir = self._get_export_dir(project.id)
        path = out_dir / f"{project.id}_export.md"

        lines = [
            f"# {project.name}",
            f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*",
            "",
        ]

        if project.transcription and project.transcription.text:
            lines += ["## Transcription", "", project.transcription.text, ""]

        for analysis in project.analyses:
            if analysis.content:
                title = analysis.analysis_type.replace("_", " ").title()
                lines += [f"## {title}", "", analysis.content, ""]

        path.write_text("\n".join(lines), encoding="utf-8")
        return path

    def export_json(self, project: Project) -> Path:
        out_dir = self._get_export_dir(project.id)
        path = out_dir / f"{project.id}_export.json"

        data = {
            "project": {
                "id": project.id,
                "name": project.name,
                "description": project.description,
                "created_at": project.created_at.isoformat(),
                "original_filename": project.original_filename,
                "audio_duration": project.audio_duration,
            },
            "transcription": None,
            "analyses": [],
        }

        if project.transcription:
            data["transcription"] = {
                "text": project.transcription.text,
                "language": project.transcription.language_detected,
                "word_count": project.transcription.word_count,
                "processing_time": project.transcription.processing_time,
                "segments": project.transcription.segments,
            }

        for analysis in project.analyses:
            if analysis.content:
                data["analyses"].append({
                    "type": analysis.analysis_type,
                    "content": analysis.content,
                    "model": analysis.model_used,
                    "created_at": analysis.created_at.isoformat(),
                })

        path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        return path

    def export_srt(self, project: Project) -> Path:
        if not project.transcription or not project.transcription.srt_file:
            raise ValueError("No SRT transcription available")
        src = Path(project.transcription.srt_file)
        out_dir = self._get_export_dir(project.id)
        dst = out_dir / f"{project.id}_export.srt"
        dst.write_bytes(src.read_bytes())
        return dst

    def export_vtt(self, project: Project) -> Path:
        if not project.transcription or not project.transcription.vtt_file:
            raise ValueError("No VTT transcription available")
        src = Path(project.transcription.vtt_file)
        out_dir = self._get_export_dir(project.id)
        dst = out_dir / f"{project.id}_export.vtt"
        dst.write_bytes(src.read_bytes())
        return dst

    def export_docx(self, project: Project) -> Path:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        title = doc.add_heading(project.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Original file: {project.original_filename or 'N/A'}")
        if project.audio_duration:
            mins = int(project.audio_duration // 60)
            secs = int(project.audio_duration % 60)
            doc.add_paragraph(f"Duration: {mins}:{secs:02d}")

        doc.add_page_break()

        if project.transcription and project.transcription.text:
            doc.add_heading("Transcription", level=1)
            doc.add_paragraph(project.transcription.text)
            doc.add_page_break()

        for analysis in project.analyses:
            if analysis.content:
                title_text = analysis.analysis_type.replace("_", " ").title()
                doc.add_heading(title_text, level=1)
                for paragraph in analysis.content.split("\n\n"):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
                doc.add_page_break()

        out_dir = self._get_export_dir(project.id)
        path = out_dir / f"{project.id}_export.docx"
        doc.save(str(path))
        return path

    def export_pdf(self, project: Project) -> Path:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib import colors

        out_dir = self._get_export_dir(project.id)
        path = out_dir / f"{project.id}_export.pdf"

        doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name="CustomBody", fontSize=10, spaceAfter=6, leading=14))
        styles.add(ParagraphStyle(name="CustomH1", fontSize=16, spaceAfter=12, spaceBefore=20,
                                  textColor=colors.HexColor("#1a1a2e"), fontName="Helvetica-Bold"))

        story = []
        story.append(Paragraph(project.name, styles["Title"]))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["CustomBody"]))
        story.append(Spacer(1, 0.3 * inch))

        if project.transcription and project.transcription.text:
            story.append(Paragraph("Transcription", styles["CustomH1"]))
            for para in project.transcription.text.split("\n\n"):
                if para.strip():
                    clean = para.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    story.append(Paragraph(clean, styles["CustomBody"]))
            story.append(PageBreak())

        for analysis in project.analyses:
            if analysis.content:
                title_text = analysis.analysis_type.replace("_", " ").title()
                story.append(Paragraph(title_text, styles["CustomH1"]))
                for para in analysis.content.split("\n\n"):
                    if para.strip():
                        clean = para.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        story.append(Paragraph(clean, styles["CustomBody"]))
                story.append(Spacer(1, 0.2 * inch))

        doc.build(story)
        return path


export_service = ExportService()
